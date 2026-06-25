from __future__ import annotations

from pathlib import Path

from docx import Document


def extract_docx_hyperlinks(path: Path) -> list[str]:
    """Return embedded hyperlink targets from a DOCX relationship store."""
    document = Document(path)
    return [
        rel.target_ref
        for rel in document.part.rels.values()
        if "hyperlink" in rel.reltype
    ]


class DocxExtractor:
    def extract_text(self, path: Path) -> str:
        document = Document(path)
        parts = [p.text for p in document.paragraphs if p.text]
        parts.extend(extract_docx_hyperlinks(path))
        return "\n".join(parts)
